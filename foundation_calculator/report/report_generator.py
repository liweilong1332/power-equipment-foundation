#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
报告生成模块 - 用于生成计算报告
"""

import os
import time
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from datetime import datetime
import matplotlib.pyplot as plt  

plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签  

plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
# PDF生成
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.platypus import PageBreak, ListFlowable, ListItem
from reportlab.graphics.shapes import Drawing, Line
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.charts.lineplots import LinePlot

# Word生成
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# 添加中文字体支持
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# 注册中文字体
def register_chinese_fonts():
    """注册中文字体以支持PDF中的中文显示"""
    try:
        # 尝试注册Windows常见中文字体
        font_paths = [
            # Windows 字体路径
            "C:/Windows/Fonts/simhei.ttf",  # 黑体
            "C:/Windows/Fonts/simsun.ttc",  # 宋体
            "C:/Windows/Fonts/simkai.ttf",  # 楷体
            "C:/Windows/Fonts/msyh.ttc",    # 微软雅黑
            "C:/Windows/Fonts/msyhbd.ttc",  # 微软雅黑粗体
            "C:/Windows/Fonts/simfang.ttf", # 仿宋
            "C:/Windows/Fonts/stkaiti.ttf", # 华文楷体
            "C:/Windows/Fonts/stzhongs.ttf",# 华文中宋
            
            # Linux 字体路径
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            
            # macOS 字体路径
            "/Library/Fonts/Arial Unicode.ttf",
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/STHeiti Medium.ttc"
        ]
        
        # 尝试注册字体，直到一个成功
        registered_fonts = []
        
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    if "simsun.ttc" in font_path:
                        pdfmetrics.registerFont(TTFont('SimSun', font_path, subfontIndex=0))
                        registered_fonts.append('SimSun')
                    elif "msyh.ttc" in font_path:
                        pdfmetrics.registerFont(TTFont('MSYaHei', font_path, subfontIndex=0))
                        registered_fonts.append('MSYaHei')
                    elif "simhei.ttf" in font_path:
                        pdfmetrics.registerFont(TTFont('SimHei', font_path))
                        registered_fonts.append('SimHei')
                    else:
                        font_name = os.path.basename(font_path).split('.')[0]
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        registered_fonts.append(font_name)
                except Exception as e:
                    print(f"注册字体 {font_path} 时出错: {str(e)}")
                    continue
        
        # 如果注册了字体，提示成功
        if registered_fonts:
            print(f"成功注册字体: {', '.join(registered_fonts)}")
        # 如果没有注册成功，使用reportlab内置字体
        else:
            print("警告：未找到中文字体，使用内置字体，中文可能显示为方块")
            
    except Exception as e:
        print(f"注册中文字体时出错: {str(e)}")
        print("继续使用默认字体，中文可能显示不正确")

# 在模块加载时注册字体
register_chinese_fonts()


def generate_report(filename, params, results):
    """
    生成计算报告
    
    参数:
    filename (str): 文件名
    params (dict): 计算参数
    results (dict): 计算结果
    
    返回:
    bool: 是否成功生成报告
    """
    try:
        # 检查是否是计算书格式
        is_calculation_book = results.get("is_calculation_book", False)
        
        # 根据文件扩展名选择生成方式
        if filename.lower().endswith('.pdf'):
            return generate_pdf_report(filename, params, results, is_calculation_book)
        elif filename.lower().endswith('.docx'):
            return generate_word_report(filename, params, results, is_calculation_book)
        else:
            # 如果没有扩展名或不支持的扩展名，默认使用PDF
            if not ('.' in os.path.basename(filename)):
                filename += '.pdf'
                return generate_pdf_report(filename, params, results, is_calculation_book)
            else:
                return False
    except Exception as e:
        print(f"生成报告时出错: {str(e)}")
        return False


def generate_pdf_report(filename, params, results, is_calculation_book=False):
    """
    生成PDF格式报告
    
    参数:
    filename (str): PDF文件名
    params (dict): 计算参数
    results (dict): 计算结果
    is_calculation_book (bool): 是否是计算书格式
    
    返回:
    bool: 是否成功生成报告
    """
    try:
        # 创建文档
        doc = SimpleDocTemplate(
            filename,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
            # 确保编码为UTF-8
            encoding='utf-8'
        )
        
        # 获取样式
        styles = getSampleStyleSheet()
        
        # 创建自定义样式，使用中文字体
        # 检查已注册的字体
        registered_fonts = pdfmetrics.getRegisteredFontNames()
        print(f"已注册的字体: {registered_fonts}")
        
        # 选择中文字体，优先级：SimHei > SimSun > MSYaHei > 其他
        chinese_font_name = None
        preferred_fonts = ['SimHei', 'SimSun', 'MSYaHei', 'simkai', 'simfang', 'stkaiti', 'stzhongs']
        
        for font in preferred_fonts:
            if font in registered_fonts:
                chinese_font_name = font
                print(f"选择字体: {chinese_font_name}")
                break
        
        # 如果没有找到合适的字体，使用默认字体
        if not chinese_font_name:
            chinese_font_name = 'Helvetica'
            print("警告：未找到中文字体，使用默认字体Helvetica，中文可能显示为方块")
        
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Title'],
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=20,
            fontName=chinese_font_name  # 使用中文字体
        )
        
        heading1_style = ParagraphStyle(
            'Heading1',
            parent=styles['Heading1'],
            fontSize=16,
            alignment=TA_LEFT,
            spaceAfter=10,
            fontName=chinese_font_name  # 使用中文字体
        )
        
        heading2_style = ParagraphStyle(
            'Heading2',
            parent=styles['Heading2'],
            fontSize=14,
            alignment=TA_LEFT,
            spaceAfter=8,
            fontName=chinese_font_name  # 使用中文字体
        )
        
        heading3_style = ParagraphStyle(
            'Heading3',
            parent=styles['Heading3'],
            fontSize=12,
            alignment=TA_LEFT,
            spaceAfter=6,
            fontName=chinese_font_name  # 使用中文字体
        )
        
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            fontName=chinese_font_name  # 使用中文字体
        )
        
        formula_style = ParagraphStyle(
            'Formula',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_CENTER,
            spaceAfter=6,
            fontName=chinese_font_name  # 使用中文字体
        )
        
        # 准备文档内容
        content = []
        
        # 添加标题
        if is_calculation_book:
            content.append(Paragraph("大块式设备基础计算书", title_style))
        else:
            content.append(Paragraph("大块式设备基础计算报告", title_style))
        content.append(Spacer(1, 0.2*inch))
        
        # 添加报告基本信息
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        content.append(Paragraph(f"生成时间: {current_time}", normal_style))
        content.append(Paragraph(f"项目名称: {params.get('name', '新计算')}", normal_style))
        content.append(Paragraph(f"项目描述: {params.get('description', '大块式设备基础计算')}", normal_style))
        content.append(Spacer(1, 0.2*inch))
        
        # 如果是计算书格式，添加目录
        if is_calculation_book:
            content.append(Paragraph("目录", heading1_style))
            content.append(Paragraph("1. 工程概况", normal_style))
            content.append(Paragraph("2. 计算依据", normal_style))
            content.append(Paragraph("3. 计算参数", normal_style))
            content.append(Paragraph("4. 静力分析", normal_style))
            content.append(Paragraph("5. 动力分析", normal_style))
            content.append(Paragraph("6. 结论与建议", normal_style))
            content.append(Paragraph("7. 附录：相关规范摘要", normal_style))
            content.append(PageBreak())
            
            # 添加工程概况
            content.append(Paragraph("1. 工程概况", heading1_style))
            project_desc = (
                "本计算书基于《动力机器基础设计标准》(GB 50040-2020)，对大块式设备基础进行静力和动力分析计算。"
                "计算内容包括基础的抗滑移稳定性、抗倾覆稳定性、地基承载力验算以及动力响应分析等。"
                f"\n\n本工程名称：{params.get('name', '新计算')}"
                f"\n工程描述：{params.get('description', '大块式设备基础计算')}"
                f"\n设计单位：-"
                f"\n设计日期：{current_time.split()[0]}"
            )
            content.append(Paragraph(project_desc, normal_style))
            content.append(Spacer(1, 0.2*inch))
            
            # 添加计算依据
            content.append(Paragraph("2. 计算依据", heading1_style))
            standards = [
                "《动力机器基础设计标准》GB 50040-2020",
                "《建筑地基基础设计规范》GB 50007-2011",
                "《混凝土结构设计规范》GB 50010-2010",
                "《建筑抗震设计规范》GB 50011-2010"
            ]
            
            standards_list = []
            for i, std in enumerate(standards):
                standards_list.append(ListItem(Paragraph(str(std), normal_style)))
                
            content.append(ListFlowable(standards_list, bulletType='bullet', start=1))
            content.append(Spacer(1, 0.2*inch))
            
            # 添加计算基本原理简介
            if is_calculation_book:
                content.append(Paragraph("计算基本原理", heading2_style))
                theory_text = (
                    "大块式设备基础设计计算主要包括两部分：静力计算和动力计算。"
                    "\n\n静力计算主要验证基础在静荷载作用下的安全性，包括抗滑移稳定性、抗倾覆稳定性和地基承载力验算；"
                    "动力计算则分析设备运行时产生的振动对基础及周围环境的影响，主要计算振动幅值是否满足要求，"
                    "以及评估共振风险。"
                    "\n\n在动力分析中，基础-设备系统简化为单自由度系统，使用频率响应法计算振动幅值，"
                    "根据《动力机器基础设计标准》(GB 50040-2020)的规定，"
                    "将计算得到的振动幅值与容许振幅进行比较，以判断设计是否满足要求。"
                )
                content.append(Paragraph(theory_text, normal_style))
                content.append(Spacer(1, 0.2*inch))
            
            content.append(PageBreak())
        
        # 添加计算参数摘要
        if is_calculation_book:
            content.append(Paragraph("3. 计算参数", heading1_style))
        else:
            content.append(Paragraph("计算参数", heading1_style))
        
        # 添加计算结果摘要
        content.append(Paragraph("计算结果摘要", heading2_style))
        
        # 判断整体安全性
        static_results = results.get("static_results", {})
        dynamic_results = results.get("dynamic_results", {})
        
        is_sliding_safe = static_results.get("sliding_stability", {}).get("is_safe", False)
        is_overturning_safe = static_results.get("overturning_stability", {}).get("is_safe", False)
        is_bearing_safe = static_results.get("bearing_capacity", {}).get("is_safe", False)
        is_dynamic_safe = dynamic_results.get("is_safe", False)
        
        overall_safety = is_sliding_safe and is_overturning_safe and is_bearing_safe and is_dynamic_safe
        
        # 安全性结果表格
        safety_data = [
            ["检验项目", "结果", "说明"],
            ["抗滑移稳定性", "满足" if is_sliding_safe else "不满足", 
             f"安全系数: {static_results.get('sliding_stability', {}).get('safety_factor', 0):.2f}"],
            ["抗倾覆稳定性", "满足" if is_overturning_safe else "不满足", 
             f"安全系数: {static_results.get('overturning_stability', {}).get('safety_factor', 0):.2f}"],
            ["地基承载力", "满足" if is_bearing_safe else "不满足", 
             f"实际压力: {static_results.get('bearing_capacity', {}).get('actual_pressure', 0):.2f} kPa"],
            ["振动响应", "满足" if is_dynamic_safe else "不满足", 
             f"振幅: {dynamic_results.get('amplitude', 0):.3f} mm"]
        ]
        
        # 使用ensure_string_data函数确保表格数据都是字符串类型
        safety_data = ensure_string_data(safety_data)
        
        safety_table = Table(safety_data, colWidths=[doc.width/3.0]*3)
        safety_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ALIGN', (1, 1), (1, -1), 'CENTER'),
            ('TEXTCOLOR', (1, 1), (1, -1), colors.green if overall_safety else colors.red)
        ]))
        
        content.append(safety_table)
        content.append(Spacer(1, 0.2*inch))
        
        # 添加基础参数
        content.append(Paragraph("基础参数", heading1_style))
        
        # 几何参数表格
        geo_data = [
            ["参数", "数值"],
            ["基础长度", f"{params.get('length', 0):.2f} m"],
            ["基础宽度", f"{params.get('width', 0):.2f} m"],
            ["基础高度", f"{params.get('height', 0):.2f} m"],
            ["埋深", f"{params.get('buried_depth', 0):.2f} m"]
        ]
        
        # 使用ensure_string_data函数确保表格数据都是字符串类型
        geo_data = ensure_string_data(geo_data)
        
        geo_table = Table(geo_data, colWidths=[doc.width/2.0]*2)
        geo_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        content.append(Paragraph("几何参数", heading2_style))
        content.append(geo_table)
        content.append(Spacer(1, 0.2*inch))
        
        # 材料参数表格
        mat_data = [
            ["参数", "数值"],
            ["混凝土强度", f"{params.get('concrete_strength', 0):.1f} MPa"],
            ["弹性模量", f"{params.get('elastic_modulus', 0):.0f} MPa"],
            ["地基承载力", f"{params.get('soil_bearing_capacity', 0):.1f} kPa"],
            ["地基系数", f"{params.get('soil_coefficient', 0):.0f} kN/m³"]
        ]
        
        # 使用ensure_string_data函数确保表格数据都是字符串类型
        mat_data = ensure_string_data(mat_data)
        
        mat_table = Table(mat_data, colWidths=[doc.width/2.0]*2)
        mat_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        content.append(Paragraph("材料参数", heading2_style))
        content.append(mat_table)
        content.append(Spacer(1, 0.2*inch))
        
        # 荷载参数表格
        load_data = [
            ["参数", "数值"],
            ["静荷载", f"{params.get('static_load', 0):.1f} kN"],
            ["动力荷载", f"{params.get('dynamic_load', 0):.1f} kN"],
            ["频率", f"{params.get('frequency', 0):.1f} Hz"],
            ["荷载偏心X", f"{params.get('load_eccentricity_x', 0):.2f} m"],
            ["荷载偏心Y", f"{params.get('load_eccentricity_y', 0):.2f} m"],
            ["设备质量", f"{params.get('equipment_mass', 0):.0f} kg"]
        ]
        
        # 使用ensure_string_data函数确保表格数据都是字符串类型
        load_data = ensure_string_data(load_data)
        
        load_table = Table(load_data, colWidths=[doc.width/2.0]*2)
        load_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        content.append(Paragraph("荷载参数", heading2_style))
        content.append(load_table)
        content.append(Spacer(1, 0.2*inch))
        
        # 添加分页
        content.append(PageBreak())
        
        # 添加静力分析结果
        if is_calculation_book:
            content.append(Paragraph("4. 静力分析", heading1_style))
        else:
            content.append(Paragraph("静力分析结果", heading1_style))
        
        # 添加计算过程说明
        content.append(Paragraph("计算过程与理论依据", heading2_style))
        
        calc_theory = (
            "静力分析基于《建筑地基基础设计规范》(GB 50007)和《动力机器基础设计标准》(GB 50040)，"
            "检验基础的抗滑移稳定性、抗倾覆稳定性和地基承载力是否满足要求。"
        )
        content.append(Paragraph(calc_theory, normal_style))
        content.append(Spacer(1, 0.1*inch))
        
        # 抗滑移稳定性
        content.append(Paragraph("抗滑移稳定性", heading2_style))
        
        # 添加计算公式和过程
        sliding_calc = (
            "抗滑移稳定性计算基于以下公式:"
            "\n\n抗滑力: F_抗滑 = mu * N"
            "\n滑动力: F_滑动 = H"
            "\n安全系数: K_滑 = F_抗滑 / F_滑动"
            "\n\n其中:"
            "\nmu - 基础底面与地基的摩擦系数, 取值 " + f"{static_results.get('sliding_stability', {}).get('friction_coefficient', 0):.2f}"
            "\nN - 基础底面法向力, 等于基础总重和设备重量之和, 为 " + f"{static_results.get('total_weight', 0):.2f} kN"
            "\nH - 水平荷载, 为 " + f"{static_results.get('horizontal_force', 0):.2f} kN"
        )
        
        # 计算过程中的数值
        anti_sliding_force = static_results.get("sliding_stability", {}).get("anti_sliding_force", 0)
        sliding_force = static_results.get("sliding_stability", {}).get("sliding_force", 0)
        sliding_factor = static_results.get("sliding_stability", {}).get("safety_factor", 0)
        sliding_required = static_results.get("sliding_stability", {}).get("required_factor", 1.3)
        
        sliding_calc += (
            f"\n\n计算过程:"
            f"\nF_抗滑 = {static_results.get('sliding_stability', {}).get('friction_coefficient', 0):.2f} * {static_results.get('total_weight', 0):.2f} = {anti_sliding_force:.2f} kN"
            f"\nF_滑动 = {sliding_force:.2f} kN"
            f"\nK_滑 = {anti_sliding_force:.2f} / {sliding_force:.2f} = {sliding_factor:.2f}"
            f"\n\n规范要求 K_滑 >= {sliding_required}, {'满足' if is_sliding_safe else '不满足'}要求。"
        )
        
        content.append(Paragraph(sliding_calc, normal_style))
        content.append(Spacer(1, 0.1*inch))
        
        # 抗倾覆稳定性
        content.append(Paragraph("抗倾覆稳定性", heading2_style))
        
        # 添加计算公式和过程
        overturning_calc = (
            "抗倾覆稳定性计算基于以下公式:"
            "\n\n抗倾覆力矩: M_抗倾 = N * L/2"
            "\n倾覆力矩: M_倾覆 = H * h + N * e"
            "\n安全系数: K_倾 = M_抗倾 / M_倾覆"
            "\n\n其中:"
            "\nN - 基础底面法向力, 等于基础总重和设备重量之和, 为 " + f"{static_results.get('total_weight', 0):.2f} kN"
            "\nL - 基础长度, 在倾覆方向上的尺寸, 为 " + f"{params.get('length', 0):.2f} m"
            "\nH - 水平荷载, 为 " + f"{static_results.get('horizontal_force', 0):.2f} kN"
            "\nh - 水平荷载作用点到基础底面的垂直距离, 为 " + f"{params.get('height', 0):.2f} m"
            "\ne - 竖向荷载偏心距, 为 " + f"{params.get('load_eccentricity_x', 0):.2f} m"
        )
        
        # 计算过程中的数值
        stabilizing_moment = static_results.get("overturning_stability", {}).get("stabilizing_moment", 0)
        overturning_moment = static_results.get("overturning_stability", {}).get("overturning_moment", 0)
        overturning_factor = static_results.get("overturning_stability", {}).get("safety_factor", 0)
        overturning_required = static_results.get("overturning_stability", {}).get("required_factor", 1.5)
        
        length = params.get('length', 0)
        height = params.get('height', 0)
        eccentricity = params.get('load_eccentricity_x', 0)
        
        overturning_calc += (
            f"\n\n计算过程:"
            f"\nM_抗倾 = {static_results.get('total_weight', 0):.2f} * {length:.2f}/2 = {stabilizing_moment:.2f} kN·m"
            f"\nM_倾覆 = {static_results.get('horizontal_force', 0):.2f} * {height:.2f} + {static_results.get('total_weight', 0):.2f} * {eccentricity:.2f} = {overturning_moment:.2f} kN·m"
            f"\nK_倾 = {stabilizing_moment:.2f} / {overturning_moment:.2f} = {overturning_factor:.2f}"
            f"\n\n规范要求 K_倾 >= {overturning_required}, {'满足' if is_overturning_safe else '不满足'}要求。"
        )
        
        content.append(Paragraph(overturning_calc, normal_style))
        content.append(Spacer(1, 0.1*inch))
        
        # 地基承载力
        content.append(Paragraph("地基承载力验算", heading2_style))
        
        # 添加计算公式和过程
        bearing_calc = (
            "地基承载力验算基于以下公式:"
            "\n\n实际地基压力: p = N / A + M / W"
            "\n其中:"
            "\nN - 基础底面法向力, 等于基础总重和设备重量之和, 为 " + f"{static_results.get('total_weight', 0):.2f} kN"
            "\nA - 基础底面积, A = L * B = " + f"{params.get('length', 0):.2f} * {params.get('width', 0):.2f} = {params.get('length', 0) * params.get('width', 0):.2f} m²"
            "\nM - 弯矩, M = N * e = " + f"{static_results.get('total_weight', 0):.2f} * {params.get('load_eccentricity_x', 0):.2f} = {static_results.get('total_weight', 0) * params.get('load_eccentricity_x', 0):.2f} kN·m"
            "\nW - 截面模量, W = L * B² / 6 = " + f"{params.get('length', 0):.2f} * {params.get('width', 0):.2f}² / 6 = {params.get('length', 0) * params.get('width', 0)**2 / 6:.2f} m³"
            "\ne - 竖向荷载偏心距, 为 " + f"{params.get('load_eccentricity_x', 0):.2f} m"
        )
        
        # 计算过程中的数值
        actual_pressure = static_results.get("bearing_capacity", {}).get("actual_pressure", 0)
        allowable_pressure = static_results.get("bearing_capacity", {}).get("allowable_pressure", 0)
        
        area = params.get('length', 0) * params.get('width', 0)
        moment = static_results.get('total_weight', 0) * params.get('load_eccentricity_x', 0)
        section_modulus = params.get('length', 0) * params.get('width', 0)**2 / 6
        
        bearing_calc += (
            f"\n\n计算过程:"
            f"\np = {static_results.get('total_weight', 0):.2f} / {area:.2f} + {moment:.2f} / {section_modulus:.2f} = {actual_pressure:.2f} kPa"
            f"\n允许地基承载力: [p] = {allowable_pressure:.2f} kPa"
            f"\n\n规范要求 p <= [p], {'满足' if is_bearing_safe else '不满足'}要求。"
        )
        
        content.append(Paragraph(bearing_calc, normal_style))
        content.append(Spacer(1, 0.2*inch))
        
        # 添加分页
        content.append(PageBreak())
        
        # 添加动力分析结果
        if is_calculation_book:
            content.append(Paragraph("5. 动力分析", heading1_style))
        else:
            content.append(Paragraph("动力分析结果", heading1_style))
        
        # 添加计算过程说明
        content.append(Paragraph("计算过程与理论依据", heading2_style))
        
        dynamic_theory = (
            "动力分析基于《动力机器基础设计标准》(GB 50040)，采用单自由度系统模型，"
            "计算基础系统的自然频率、振动响应以及传递给周围环境的振动大小，"
            "检验振动幅值是否满足要求，并评估共振风险。"
        )
        content.append(Paragraph(dynamic_theory, normal_style))
        content.append(Spacer(1, 0.1*inch))
        
        # 自然频率计算
        content.append(Paragraph("自然频率计算", heading2_style))
        
        # 添加计算公式和过程
        freq_calc = (
            "基础系统自然频率计算基于以下公式:"
            "\n\n自然频率: f_n = (1/2pi) * sqrt(K/M)"
            "\n其中:"
            "\nK - 系统刚度, K = A * k = " + f"{params.get('length', 0) * params.get('width', 0):.2f} * {params.get('soil_coefficient', 0):.0f} = {params.get('length', 0) * params.get('width', 0) * params.get('soil_coefficient', 0):.0f} kN/m"
            "\nM - 系统质量, M = m_基础 + m_设备 = " + f"{results.get('foundation_mass', 0):.0f} + {params.get('equipment_mass', 0):.0f} = {results.get('total_mass', 0):.0f} kg"
            "\nA - 基础底面积, A = " + f"{params.get('length', 0) * params.get('width', 0):.2f} m²"
            "\nk - 地基系数, k = " + f"{params.get('soil_coefficient', 0):.0f} kN/m³"
        )
        
        # 计算过程中的数值
        natural_frequency = dynamic_results.get("natural_frequency", 0)
        stiffness = params.get('length', 0) * params.get('width', 0) * params.get('soil_coefficient', 0)
        
        freq_calc += (
            f"\n\n计算过程:"
            f"\nf_n = (1/2pi) * sqrt({stiffness:.0f}/{results.get('total_mass', 0):.0f}) = {natural_frequency:.2f} Hz"
        )
        
        content.append(Paragraph(freq_calc, normal_style))
        content.append(Spacer(1, 0.1*inch))
        
        # 振动响应计算
        content.append(Paragraph("振动响应计算", heading2_style))
        
        # 添加计算公式和过程
        response_calc = (
            "振动响应计算基于以下公式："
            "\n\n频率比: r = f/f_n"
            "\n振动传递率: T = 1/sqrt[(1-r^2)^2 + (2*zeta*r)^2]"
            "\n振幅: A = F*T/(K)"
            "\n\n其中:"
            "\nf - 激励频率, f = " + f"{dynamic_results.get('frequency_ratio', 0) * natural_frequency:.2f} Hz"
            "\nf_n - 系统自然频率, f_n = " + f"{natural_frequency:.2f} Hz"
            "\nzeta - 阻尼比, zeta = " + f"{params.get('damping_ratio', 0):.2f}"
            "\nF - 动力荷载, F = " + f"{params.get('dynamic_load', 0):.1f} kN"
            "\nK - 系统刚度, K = " + f"{stiffness:.0f} kN/m"
        )
        
        # 计算过程中的数值
        frequency_ratio = dynamic_results.get("frequency_ratio", 0)
        transmissibility = dynamic_results.get("transmissibility", 0)
        amplitude = dynamic_results.get("amplitude", 0)
        allowable_amplitude = dynamic_results.get("allowable_amplitude", 0)
        
        response_calc += (
            f"\n\n计算过程:"
            f"\nr = {dynamic_results.get('frequency_ratio', 0) * natural_frequency:.2f}/{natural_frequency:.2f} = {frequency_ratio:.2f}"
            f"\nT = 1/sqrt[(1-{frequency_ratio:.2f}^2)^2 + (2*{params.get('damping_ratio', 0):.2f}*{frequency_ratio:.2f})^2] = {transmissibility:.2f}"
            f"\nA = {params.get('dynamic_load', 0):.1f}*{transmissibility:.2f}/{stiffness:.0f} = {amplitude:.3f} m = {amplitude*1000:.3f} mm"
            f"\n允许振幅: [A] = {allowable_amplitude:.3f} mm"
            f"\n\n规范要求 A <= [A], {'满足' if is_dynamic_safe else '不满足'}要求。"
        )
        
        content.append(Paragraph(response_calc, normal_style))
        content.append(Spacer(1, 0.1*inch))
        
        # 共振风险评估
        content.append(Paragraph("共振风险评估", heading2_style))
        
        # 添加评估过程
        resonance_risk = dynamic_results.get("resonance_risk", False)
        
        resonance_calc = (
            "共振风险评估基于频率比 r:\n"
            "* 当 0.8 < r < 1.2 时, 系统处于共振区, 存在共振风险\n"
            "* 当 r >= 1.4 时, 系统运行在超共振区, 振动传递率较低\n"
            "* 当 r <= 0.7 时, 系统运行在次共振区, 振动传递率接近1\n"
            f"\n当前频率比 r = {frequency_ratio:.2f}, "
        )
        
        if frequency_ratio < 0.8:
            resonance_calc += f"系统运行在次共振区, 振动传递率接近1, {'不存在' if not resonance_risk else '存在'}共振风险。"
        elif frequency_ratio > 1.2:
            resonance_calc += f"系统运行在超共振区, 振动传递率较低, {'不存在' if not resonance_risk else '存在'}共振风险。"
        else:
            resonance_calc += f"系统处于共振区, {'存在' if resonance_risk else '不存在'}共振风险。"
        
        content.append(Paragraph(resonance_calc, normal_style))
        content.append(Spacer(1, 0.2*inch))
        
        # 频率响应曲线图
        content.append(Paragraph("频率响应曲线", heading2_style))
        
        # 创建频率响应曲线图
        if "response_curve" in results:
            frequencies = results["response_curve"].get("frequencies", [])
            amplitudes = results["response_curve"].get("amplitudes", [])
            
            # 确保数据是数值型的
            try:
                # 修复：确保所有数据都是可绘制的数值
                frequencies = [float(f) if isinstance(f, (int, float, str)) else 0.0 for f in frequencies]
                amplitudes = [float(a) if isinstance(a, (int, float, str)) else 0.0 for a in amplitudes]
                
                # 使用matplotlib创建图像
                plt.figure(figsize=(8, 4))
                plt.plot(frequencies, amplitudes, 'b-')
                plt.axvline(x=float(natural_frequency), color='r', linestyle='--', alpha=0.7)
                plt.axhline(y=float(allowable_amplitude), color='g', linestyle='-.', alpha=0.7)
                plt.grid(True)
                plt.title("频率响应曲线")
                plt.xlabel("频率 (Hz)")
                plt.ylabel("振幅 (mm)")
                plt.legend(["响应曲线", "自然频率", "容许振幅"])
                
                # 保存图像到内存
                img_buffer = BytesIO()
                plt.savefig(img_buffer, format='png', dpi=100)
                img_buffer.seek(0)
                plt.close()
                
                # 添加图像到PDF
                img = Image(img_buffer, width=6*inch, height=3*inch)
                content.append(img)
            except Exception as e:
                # 如果绘图失败，添加错误信息
                content.append(Paragraph(f"生成频率响应曲线失败: {str(e)}", normal_style))
        
        content.append(Spacer(1, 0.2*inch))
        
        # 结论
        if is_calculation_book:
            content.append(Paragraph("6. 结论与建议", heading1_style))
        else:
            content.append(Paragraph("结论", heading1_style))
        
        # 根据计算结果生成结论
        conclusion_text = "基于以上计算结果，该大块式设备基础设计"
        
        if overall_safety:
            conclusion_text += "满足所有安全要求，可以实施。"
        else:
            conclusion_text += "不满足所有安全要求，建议进行如下修改："
            
        content.append(Paragraph(conclusion_text, normal_style))
        
        # 如果不满足要求，提供修改建议
        if not overall_safety:
            suggestions = []
            
            if not is_sliding_safe:
                suggestions.append("增加基础尺寸或埋深以提高抗滑移稳定性")
                
            if not is_overturning_safe:
                suggestions.append("增加基础尺寸或调整荷载偏心距以提高抗倾覆稳定性")
                
            if not is_bearing_safe:
                suggestions.append("增加基础尺寸以降低地基压力，或选择承载力更高的地基")
                
            if not is_dynamic_safe:
                suggestions.append("调整基础尺寸或质量以改变系统固有频率，避开共振区域")
            
            # 确保建议是字符串类型
            suggestions = ensure_string_data(suggestions)
            
            # 添加建议列表
            suggestion_list = []
            for i, s in enumerate(suggestions):
                suggestion_list.append(ListItem(Paragraph(s, normal_style)))
                
            content.append(ListFlowable(suggestion_list, bulletType='bullet', start=1))
        
        # 如果是计算书，添加附录
        if is_calculation_book:
            content.append(PageBreak())
            content.append(Paragraph("7. 附录：相关规范摘要", heading1_style))
            
            # 添加规范摘要
            content.append(Paragraph("《动力机器基础设计标准》(GB 50040-2020) 相关条文", heading2_style))
            
            code_extract = (
                "5.1.4 大块式基础的抗倾覆稳定性和抗滑移稳定性应按下列规定验算："
                "\n1. 基础的抗倾覆稳定性应按下式验算：M抗 / M倾 ≥ [K倾]"
                "\n其中：M抗为基础抗倾覆力矩；M倾为使基础倾覆的力矩；[K倾]为抗倾覆稳定性的安全系数，不应小于1.5。"
                "\n\n2. 基础的抗滑移稳定性应按下式验算：F抗 / F滑 ≥ [K滑]"
                "\n其中：F抗为基础抗滑移的力；F滑为使基础滑移的力；[K滑]为抗滑移稳定性的安全系数，不应小于1.3。"
                "\n\n5.2.3 基础底面的平均压力和最大压力应满足下列要求："
                "\n1. 基础底面平均压力应按下式计算：p = N / A"
                "\n其中：N为基础底面法向力；A为基础底面积。"
                "\n\n2. 基础底面平均压力不应大于地基土的承载力特征值。"
                "\n\n7.2.6 对于设计不平衡质量的往复或旋转运动机器，其基础振幅的验算，应按本标准第7.2.2条的规定进行。"
                "\n\n7.2.2 基础振幅验算应符合下列规定："
                "\n1. 当采用弹性分析时，垂直、水平摆动和水平平动方向的振幅均应满足下列要求：A ≤ [A]"
                "\n其中，A为计算振幅；[A]为容许振幅。"
                "\n\n2. 容许振幅[A]应按本标准第7.2.3~7.2.5条的规定采用。"
                "\n\n7.2.3 机器功率不大于100kW并且转速小于3000r/min时，容许振幅不应大于下表规定的值。"
            )
            content.append(Paragraph(code_extract, normal_style))
            
            # 添加容许振幅表格
            content.append(Paragraph("容许振幅表 (摘自GB 50040-2020)", heading3_style))
            amplitude_data = [
                ["机器转速 (r/min)", "垂直振幅 (mm)", "水平振幅 (mm)"],
                ["< 500", "0.2", "0.15"],
                ["500 ~ 750", "0.15", "0.12"],
                ["750 ~ 1000", "0.12", "0.10"],
                ["1000 ~ 1500", "0.10", "0.08"],
                ["> 1500", "0.08", "0.06"]
            ]
            
            # 使用ensure_string_data函数确保表格数据都是字符串类型
            amplitude_data = ensure_string_data(amplitude_data)
            
            amplitude_table = Table(amplitude_data, colWidths=[doc.width/3.0]*3)
            amplitude_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            content.append(amplitude_table)
            content.append(Spacer(1, 0.2*inch))
        
        # 构建PDF文档
        doc.build(content)
        
        return True
    except Exception as e:
        print(f"生成PDF报告时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def generate_word_report(filename, params, results, is_calculation_book=False):
    """
    生成Word格式报告
    
    参数:
    filename (str): Word文件名
    params (dict): 计算参数
    results (dict): 计算结果
    is_calculation_book (bool): 是否是计算书格式
    
    返回:
    bool: 是否成功生成报告
    """
    try:
        # 创建文档
        doc = Document()
        
        # 设置中文字体
        # 设置默认字体
        try:
            doc.styles['Normal'].font.name = 'SimSun'  # 宋体
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            print("成功设置文档默认字体为宋体(SimSun)")
        except Exception as e:
            print(f"设置文档默认字体时出错: {str(e)}")
            print("将使用默认字体")
        
        # 设置标题字体
        try:
            # 尝试多种中文字体，优先使用黑体
            heading_fonts = ['SimHei', 'Microsoft YaHei', 'SimSun', 'KaiTi']
            heading_font = 'SimHei'  # 默认使用黑体
            
            doc.styles['Heading1'].font.name = heading_font
            doc.styles['Heading1']._element.rPr.rFonts.set(qn('w:eastAsia'), heading_font)
            
            doc.styles['Heading2'].font.name = heading_font
            doc.styles['Heading2']._element.rPr.rFonts.set(qn('w:eastAsia'), heading_font)
            
            doc.styles['Heading3'].font.name = heading_font
            doc.styles['Heading3']._element.rPr.rFonts.set(qn('w:eastAsia'), heading_font)
            
            print(f"成功设置标题字体为{heading_font}")
        except Exception as e:
            print(f"设置标题字体时出错: {str(e)}")
            print("将使用默认字体")
        
        # 添加标题
        if is_calculation_book:
            doc.add_heading("大块式设备基础计算书", level=0)
        else:
            doc.add_heading("大块式设备基础计算报告", level=0)
        
        # 添加报告基本信息
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"生成时间: {current_time}")
        doc.add_paragraph(f"项目名称: {params.get('name', '新计算')}")
        doc.add_paragraph(f"项目描述: {params.get('description', '大块式设备基础计算')}")
        
        # 如果是计算书格式，添加目录
        if is_calculation_book:
            doc.add_heading("目录", level=1)
            
            # 自动目录（需要手动更新）
            doc.add_paragraph("1. 工程概况")
            doc.add_paragraph("2. 计算依据")
            doc.add_paragraph("3. 计算参数")
            doc.add_paragraph("4. 静力分析")
            doc.add_paragraph("5. 动力分析")
            doc.add_paragraph("6. 结论与建议")
            doc.add_paragraph("7. 附录：相关规范摘要")
            
            doc.add_page_break()
            
            # 添加工程概况
            doc.add_heading("1. 工程概况", level=1)
            project_desc = (
                "本计算书基于《动力机器基础设计标准》(GB 50040-2020)，对大块式设备基础进行静力和动力分析计算。"
                "计算内容包括基础的抗滑移稳定性、抗倾覆稳定性、地基承载力验算以及动力响应分析等。"
                f"\n\n本工程名称：{params.get('name', '新计算')}"
                f"\n工程描述：{params.get('description', '大块式设备基础计算')}"
                f"\n设计单位：-"
                f"\n设计日期：{current_time.split()[0]}"
            )
            doc.add_paragraph(project_desc)
            
            # 添加计算依据
            doc.add_heading("2. 计算依据", level=1)
            standards = [
                "《动力机器基础设计标准》GB 50040-2020",
                "《建筑地基基础设计规范》GB 50007-2011",
                "《混凝土结构设计规范》GB 50010-2010",
                "《建筑抗震设计规范》GB 50011-2010"
            ]
            
            for std in standards:
                doc.add_paragraph(str(std), style="List Bullet")
            
            # 添加计算基本原理简介
            doc.add_heading("计算基本原理", level=2)
            theory_text = (
                "大块式设备基础设计计算主要包括两部分：静力计算和动力计算。"
                "\n\n静力计算主要验证基础在静荷载作用下的安全性，包括抗滑移稳定性、抗倾覆稳定性和地基承载力验算；"
                "动力计算则分析设备运行时产生的振动对基础及周围环境的影响，主要计算振动幅值是否满足要求，"
                "以及评估共振风险。"
                "\n\n在动力分析中，基础-设备系统简化为单自由度系统，使用频率响应法计算振动幅值，"
                "根据《动力机器基础设计标准》(GB 50040-2020)的规定，"
                "将计算得到的振动幅值与容许振幅进行比较，以判断设计是否满足要求。"
            )
            doc.add_paragraph(theory_text)
            
            doc.add_page_break()
        
        # 添加计算参数摘要
        if is_calculation_book:
            doc.add_heading("3. 计算参数", level=1)
        else:
            doc.add_heading("计算参数", level=1)
        
        # 添加计算结果摘要
        doc.add_heading("计算结果摘要", level=2)
        
        # 判断整体安全性
        static_results = results.get("static_results", {})
        dynamic_results = results.get("dynamic_results", {})
        
        is_sliding_safe = static_results.get("sliding_stability", {}).get("is_safe", False)
        is_overturning_safe = static_results.get("overturning_stability", {}).get("is_safe", False)
        is_bearing_safe = static_results.get("bearing_capacity", {}).get("is_safe", False)
        is_dynamic_safe = dynamic_results.get("is_safe", False)
        
        overall_safety = is_sliding_safe and is_overturning_safe and is_bearing_safe and is_dynamic_safe
        
        # 安全性结果表格
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = "检验项目"
        header_cells[1].text = "结果"
        header_cells[2].text = "说明"
        
        # 添加数据行
        data = [
            ["抗滑移稳定性", "满足" if is_sliding_safe else "不满足", 
             f"安全系数: {static_results.get('sliding_stability', {}).get('safety_factor', 0):.2f}"],
            ["抗倾覆稳定性", "满足" if is_overturning_safe else "不满足", 
             f"安全系数: {static_results.get('overturning_stability', {}).get('safety_factor', 0):.2f}"],
            ["地基承载力", "满足" if is_bearing_safe else "不满足", 
             f"实际压力: {static_results.get('bearing_capacity', {}).get('actual_pressure', 0):.2f} kPa"],
            ["振动响应", "满足" if is_dynamic_safe else "不满足", 
             f"振幅: {dynamic_results.get('amplitude', 0):.3f} mm"]
        ]
        
        # 使用ensure_string_data函数确保表格数据都是字符串类型
        data = ensure_string_data(data)
        
        for item in data:
            row_cells = table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
            row_cells[2].text = item[2]
        
        doc.add_paragraph()
        
        # 添加基础参数
        doc.add_heading("基础参数", level=2)
        
        # 几何参数
        doc.add_heading("几何参数", level=3)
        
        geo_table = doc.add_table(rows=1, cols=2)
        geo_table.style = 'Table Grid'
        
        header_cells = geo_table.rows[0].cells
        header_cells[0].text = "参数"
        header_cells[1].text = "数值"
        
        geo_data = [
            ["基础长度", f"{params.get('length', 0):.2f} m"],
            ["基础宽度", f"{params.get('width', 0):.2f} m"],
            ["基础高度", f"{params.get('height', 0):.2f} m"],
            ["埋深", f"{params.get('buried_depth', 0):.2f} m"]
        ]
        
        # 使用ensure_string_data函数确保表格数据都是字符串类型
        geo_data = ensure_string_data(geo_data)
        
        for item in geo_data:
            row_cells = geo_table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
        
        doc.add_paragraph()
        
        # 材料参数
        doc.add_heading("材料参数", level=3)
        
        mat_table = doc.add_table(rows=1, cols=2)
        mat_table.style = 'Table Grid'
        
        header_cells = mat_table.rows[0].cells
        header_cells[0].text = "参数"
        header_cells[1].text = "数值"
        
        mat_data = [
            ["混凝土强度", f"{params.get('concrete_strength', 0):.1f} MPa"],
            ["弹性模量", f"{params.get('elastic_modulus', 0):.0f} MPa"],
            ["地基承载力", f"{params.get('soil_bearing_capacity', 0):.1f} kPa"],
            ["地基系数", f"{params.get('soil_coefficient', 0):.0f} kN/m³"]
        ]
        
        # 使用ensure_string_data函数确保表格数据都是字符串类型
        mat_data = ensure_string_data(mat_data)
        
        for item in mat_data:
            row_cells = mat_table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
        
        doc.add_paragraph()
        
        # 荷载参数
        doc.add_heading("荷载参数", level=3)
        
        load_table = doc.add_table(rows=1, cols=2)
        load_table.style = 'Table Grid'
        
        header_cells = load_table.rows[0].cells
        header_cells[0].text = "参数"
        header_cells[1].text = "数值"
        
        load_data = [
            ["静荷载", f"{params.get('static_load', 0):.1f} kN"],
            ["动力荷载", f"{params.get('dynamic_load', 0):.1f} kN"],
            ["频率", f"{params.get('frequency', 0):.1f} Hz"],
            ["荷载偏心X", f"{params.get('load_eccentricity_x', 0):.2f} m"],
            ["荷载偏心Y", f"{params.get('load_eccentricity_y', 0):.2f} m"],
            ["设备质量", f"{params.get('equipment_mass', 0):.0f} kg"]
        ]
        
        # 使用ensure_string_data函数确保表格数据都是字符串类型
        load_data = ensure_string_data(load_data)
        
        for item in load_data:
            row_cells = load_table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
        
        doc.add_paragraph()
        
        # 保存文档
        doc.save(filename)
        return True
        
    except Exception as e:
        print(f"生成Word报告时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# 确保在生成报告中的所有表格数据都是字符串类型
def ensure_string_data(data_list):
    """确保列表中的所有数据项都是字符串类型"""
    if isinstance(data_list, list):
        return [ensure_string_data(item) for item in data_list]
    elif isinstance(data_list, dict):
        return {key: ensure_string_data(value) for key, value in data_list.items()}
    else:
        return str(data_list)